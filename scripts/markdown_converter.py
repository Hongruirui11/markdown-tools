#!/usr/bin/env python3
"""
Markdown Converter - Convert Markdown to DOCX, HTML, and TXT formats
"""

import re
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import markdown
from bs4 import BeautifulSoup

def convert_markdown_to_html(input_file, output_file=None):
    """
    Convert Markdown to HTML
    
    Args:
        input_file: Path to the input Markdown file
        output_file: Path to the output HTML file. If not provided, creates input_file.html
        
    Returns:
        Path to the generated HTML file
    """
    try:
        # Read Markdown content
        with open(input_file, 'r', encoding='utf-8') as f:
            markdown_content = f.read()
        
        # Convert to HTML
        html_content = markdown.markdown(markdown_content, extensions=[
            'fenced_code',
            'tables',
            'nl2br',
            'sane_lists'
        ])
        
        # Generate output filename if not provided
        if not output_file:
            output_file = os.path.splitext(input_file)[0] + '.html'
        
        # Write HTML content to file
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        return f"HTML file saved to {output_file}"
    except Exception as e:
        raise Exception(f"Error converting to HTML: {str(e)}")

def convert_markdown_to_pdf(input_file, output_file=None):
    """Convert Markdown to PDF"""
    raise NotImplementedError("PDF conversion not implemented yet")

def convert_markdown_to_plain_text(input_file, output_file=None):
    """
    Convert Markdown to Plain Text
    
    Args:
        input_file: Path to the input Markdown file
        output_file: Path to the output plain text file. If not provided, creates input_file.txt
        
    Returns:
        Path to the generated plain text file
    """
    try:
        # Read Markdown content
        with open(input_file, 'r', encoding='utf-8') as f:
            markdown_content = f.read()
        
        # Convert to HTML first
        html_content = markdown.markdown(markdown_content, extensions=[
            'fenced_code',
            'tables',
            'nl2br',
            'sane_lists'
        ])
        
        # Parse HTML and extract plain text
        soup = BeautifulSoup(html_content, 'html.parser')
        plain_text = soup.get_text()
        
        # Generate output filename if not provided
        if not output_file:
            output_file = os.path.splitext(input_file)[0] + '.txt'
        
        # Write plain text to file
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(plain_text)
        
        return f"Plain text file saved to {output_file}"
    except Exception as e:
        raise Exception(f"Error converting to plain text: {str(e)}")


def convert_markdown_to_docx(input_file, output_file=None):
    """
    Convert Markdown to DOCX
    
    Args:
        input_file: Path to the input Markdown file
        output_file: Path to the output DOCX file. If not provided, creates input_file.docx
        
    Returns:
        Path to the generated DOCX file
    """
    try:
        # Generate output filename if not provided
        if not output_file:
            output_file = os.path.splitext(input_file)[0] + '.docx'
        
        converter = MarkdownToDocxConverter(input_file, output_file)
        converter.convert_file()
        
        return f"DOCX file saved to {output_file}"
    except Exception as e:
        raise Exception(f"Error converting to DOCX: {str(e)}")


class MarkdownToDocxConverter:
    """A class to convert Markdown files to Word documents with preserved formatting."""

    # Default styles
    DEFAULT_STYLES = {
        'document': {
            'font_name': '宋体',
            'font_size': 11,
            'margins': {
                'top': 1.0,
                'bottom': 1.0,
                'left': 1.0,
                'right': 1.0
            }
        },
        'h1': {
            'font_name': '宋体',
            'font_size': 16,
            'color': (0, 0, 0),
            'bold': True
        },
        'h2': {
            'font_name': '宋体',
            'font_size': 14,
            'color': (0, 0, 0),
            'bold': True
        },
        'h3': {
            'font_name': '宋体',
            'font_size': 12,
            'color': (0, 0, 0),
            'bold': True
        },
        'h4': {
            'font_name': '宋体',
            'font_size': 11,
            'color': (0, 0, 0),
            'bold': True
        },
        'h5': {
            'font_name': '宋体',
            'font_size': 11,
            'color': (0, 0, 0),
            'bold': True
        },
        'h6': {
            'font_name': '宋体',
            'font_size': 11,
            'color': (0, 0, 0),
            'bold': True
        },
        'paragraph': {
            'font_name': '宋体',
            'font_size': 11,
            'color': (0, 0, 0)
        },
        'code': {
            'font_name': 'Courier New',
            'font_size': 10,
            'color': (169, 169, 169)
        },
        'strong': {
            'font_name': '宋体',
            'bold': True,
            'color': (0, 0, 0)
        },
        'emphasis': {
            'font_name': '宋体',
            'italic': True,
            'color': (0, 0, 0)
        },
        'table_header': {
            'font_name': '宋体',
            'bold': True,
            'color': (0, 0, 0)
        },
        'table_cell': {
            'font_name': '宋体',
            'color': (0, 0, 0)
        }
    }

    def __init__(self, input_file, output_file, template_path=None):
        """
        Initialize the converter.
        
        Args:
            input_file: Path to the input Markdown file
            output_file: Path to the output Word document
            template_path: Path to the Word template file (optional)
        """
        self.input_file = input_file
        self.output_file = output_file
        
        # Initialize styles
        self.styles = self.DEFAULT_STYLES.copy()
        
        # Default template path - using relative path based on script location
        script_dir = os.path.dirname(os.path.abspath(__file__))
        default_template = os.path.join(script_dir, '../assets/word-template.docx')
        
        # Check if template exists
        if template_path and os.path.exists(template_path):
            self.template_path = template_path
        elif os.path.exists(default_template):
            self.template_path = default_template
        else:
            self.template_path = None
            
        # Create document from template or new document
        self.template_loaded = False
        if self.template_path:
            try:
                # Try to load the template directly
                self.doc = Document(self.template_path)
                # Clear all content from the template while preserving styles
                self._clear_template_content()
                self.template_loaded = True
            except Exception as e:
                # If loading fails (e.g., .dotx file), create a new document and apply styles manually
                print(f"Warning: Failed to load template {self.template_path}: {str(e)}")
                self.doc = Document()
        else:
            self.doc = Document()
        
        self._setup_default_document()

    def convert_file(self):
        """
        Convert the Markdown file to a Word document.
        """
        try:
            with open(self.input_file, 'r', encoding='utf-8') as f:
                markdown_text = f.read()
            
            self.convert_string(markdown_text, self.output_file)
            return True
        except Exception as e:
            print(f"Error during conversion: {str(e)}")
            return False

    def convert_string(self, markdown_text, output_file):
        """Convert a Markdown string to a Word document.
        
        Args:
            markdown_text: Markdown formatted text
            output_file: Path where the output Word document will be saved
        """
        # Process full-width space placeholders first
        text = re.sub(r'\[FULLWIDTH[ _]?SPACES:(\d+)\]', lambda m: '　' * int(m.group(1)), markdown_text, flags=re.IGNORECASE)
        
        # Convert Markdown to HTML with nl2br extension for proper table handling
        html = markdown.markdown(text, extensions=['extra', 'nl2br'])
        # Fix invalid HTML tags - comprehensive fix for br tags
        html = re.sub(r'</br>', '', html, flags=re.IGNORECASE)
        html = re.sub(r'<br\s*>', '<br/>', html, flags=re.IGNORECASE)
        
        # Parse HTML with BeautifulSoup
        soup = BeautifulSoup(html, 'html.parser')
        
        # 获取所有子元素
        children = list(soup.children)
        i = 0
        
        # 处理所有元素
        while i < len(children):
            current = children[i]
            
            # 处理当前元素
            self._process_element(current)
            i += 1
        
        # Remove extra empty paragraphs at the end
        while len(self.doc.paragraphs) > 0 and not self.doc.paragraphs[-1].text.strip():
            p = self.doc.paragraphs[-1]._element
            p.getparent().remove(p)
        
        # Post-processing loop removed: Indentation is now handled element-wise

        
        # Save the document
        self.doc.save(output_file)

    def _setup_default_document(self):
        """Configure default document margins and sections."""
        section = self.doc.sections[0]
        margins = self.styles['document']['margins']
        section.top_margin = Inches(margins['top'])
        section.bottom_margin = Inches(margins['bottom'])
        section.left_margin = Inches(margins['left'])
        section.right_margin = Inches(margins['right'])

    def _clear_template_content(self):
        """Clear all content from the template document while preserving styles and section settings."""
        # Clear all paragraphs
        while self.doc.paragraphs:
            p = self.doc.paragraphs[0]._element
            p.getparent().remove(p)
            p._p = p._element = None
        
        # Clear all tables
        while self.doc.tables:
            tbl = self.doc.tables[0]._element
            tbl.getparent().remove(tbl)
            tbl._tbl = tbl._element = None
        
        # Remove all sections except the first one
        if len(self.doc.sections) > 1:
            for section in self.doc.sections[1:]:
                section._element.getparent().remove(section._element)

    def _apply_style(self, run, style_dict):
        """Apply style to a run."""
        try:
            if 'font_name' in style_dict:
                run.font.name = style_dict['font_name']
                # Set east Asia font for Chinese characters
                run._element.rPr.rFonts.set(qn('w:eastAsia'), style_dict['font_name'])
            
            if 'font_size' in style_dict:
                run.font.size = Pt(style_dict['font_size'])
            
            if 'color' in style_dict:
                run.font.color.rgb = RGBColor(*style_dict['color'])
            
            if 'bold' in style_dict:
                run.bold = style_dict['bold']
            
            if 'italic' in style_dict:
                run.italic = style_dict['italic']
                
        except Exception as e:
            print(f"Error applying style: {str(e)}")
    
    def _parse_css_length(self, length_str, default_unit='pt'):
        """Parse a CSS length value and convert it to Pt (points).
        
        Args:
            length_str: CSS length string (e.g., '14pt', '16px', '2em')
            default_unit: Default unit if not specified (default: 'pt')
            
        Returns:
            Pt object representing the length in points
        """
        if not length_str:
            return None
        
        # Extract numeric value and unit
        match = re.match(r'(\d+(?:\.\d+)?)\s*(pt|px|em|rem)?', length_str)
        if not match:
            return None
        
        value = float(match.group(1))
        unit = match.group(2) or default_unit
        
        # Convert to points
        if unit == 'px':
            return Pt(value * 0.75)  # 1px ~ 0.75pt
        elif unit == 'em' or unit == 'rem':
            # Assume 1em = 11pt (default font size)
            return Pt(value * 11)
        else:
            return Pt(value)
    
    def _add_inline_content(self, paragraph, text, style_key='paragraph', bold=False, italic=False, font_name=None, font_size=None):
        """Add inline content to a paragraph with consistent styling.
        
        Args:
            paragraph: Word paragraph object to add content to
            text: Text content to add
            style_key: Key for default style from self.styles (default: 'paragraph')
            bold: Whether to make the text bold (default: False)
            italic: Whether to make the text italic (default: False)
            font_name: Optional font name to apply
            font_size: Optional font size to apply (Pt object)
            
        Returns:
            The created run object
        """
        if not text.strip():
            return
        
        # Create run and apply basic style
        run = paragraph.add_run(text)
        self._apply_style(run, self.styles[style_key])
        
        # Apply formatting
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        
        # Apply div-level font styles if specified
        if font_name:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        if font_size:
            run.font.size = font_size
        
        return run

    def _add_heading(self, text, level):
        """Add a heading with the corresponding Word style from the template."""
        
        # Treatment for H5 and H6: Use Body Text (Normal) style
        if level >= 5:
            # Create a regular paragraph
            p = self.doc.add_paragraph()
            # Add text run
            run = p.add_run(text)
            
            # Apply normal paragraph style
            self._apply_style(run, self.styles['paragraph'])
            
            # Ensure it uses the default paragraph style of the document (Normal/正文)
            # We explicitly don't assign "Heading X" style here
            
            # Explicitly remove indentation for H5/H6 paragraphs to differentiate from body text
            p.paragraph_format.first_line_indent = Pt(0)
            
            return p

        # Map Markdown heading levels to Word template styles
        template_style_map = {
            1: "标题 1",   # h1 -> Word template "标题 1" style
            2: "标题 2",   # h2 -> Word template "标题 2" style
            3: "标题 3",   # h3 -> Word template "标题 3" style
            4: "标题 4",   # h4 -> Word template "标题 4" style
        }
        
        # Map Markdown heading levels to Word built-in styles (fallback)
        builtin_style_map = {
            1: "Heading 1",   # h1 -> Word "Heading 1" style
            2: "Heading 2",   # h2 -> Word "Heading 2" style
            3: "Heading 3",   # h3 -> Word "Heading 3" style
            4: "Heading 4",   # h4 -> Word "Heading 4" style
        }
        
        # Try to use template style if available, otherwise use built-in style
        word_style = template_style_map.get(level, f"标题 {level}")
        
        # Add paragraph and apply style
        paragraph = self.doc.add_paragraph(text)
        
        try:
            # Try to apply template style
            paragraph.style = word_style
        except Exception:
            # Fallback to built-in style if template style not found
            builtin_style = builtin_style_map.get(level, f"Heading {level}")
            try:
                paragraph.style = builtin_style
            except:
                # If even builtin fails, fall back to Normal
                paragraph.style = "Normal"
        
        # Only apply custom styles if template is not loaded (template styles should take precedence)
        if not self.template_loaded:
            style_key = f'h{level}'
            if style_key in self.styles:
                # Apply custom style to all runs in the paragraph
                for run in paragraph.runs:
                    self._apply_style(run, self.styles[style_key])
        
        return paragraph

    def _process_element(self, element, context=None):
        """Process an HTML element and convert it to Word format."""
        if context is None:
            context = {}

        if isinstance(element, str):
            if element.strip():
                p = self.doc.add_paragraph()
                run = p.add_run(element)
                self._apply_style(run, self.styles['paragraph'])
                # Apply context styles
                if 'font_name' in context:
                    run.font.name = context['font_name']
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), context['font_name'])
                if 'font_size' in context:
                    run.font.size = context['font_size']
                # Apply text alignment from context
                if 'text_align' in context:
                    text_align = context['text_align'].lower()
                    if text_align == 'center':
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    elif text_align == 'right':
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    elif text_align == 'left':
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                # Apply text indent from context
                if 'text_indent' in context:
                    p.paragraph_format.first_line_indent = context['text_indent']
                else:
                    # Default indentation for body text (2 chars / ~21pt)
                    p.paragraph_format.first_line_indent = Pt(21)
            return

        if element.name is None:
            return

        # Parse element styles and update context
        current_context = context.copy()
        if element.name == 'div' or 'style' in element.attrs:
            style_str = element.attrs.get('style', '')
            # Split style declarations
            styles = style_str.split(';')
            for style in styles:
                style = style.strip()
                if not style:
                    continue
                # Split property and value
                if ':' in style:
                    prop, val = style.split(':', 1)
                    prop = prop.strip().lower()
                    val = val.strip()
                    
                    # Handle font-family
                    if prop == 'font-family':
                        # Remove quotes if present
                        val = val.strip('"\'')
                        # Get first font in the list
                        current_context['font_name'] = val.split(',')[0].strip()
                    # Handle font-size
                    elif prop == 'font-size':
                        current_context['font_size'] = self._parse_css_length(val)
                    # Handle text-align
                    elif prop == 'text-align':
                        current_context['text_align'] = val
                    # Handle text-indent
                    elif prop == 'text-indent':
                        current_context['text_indent'] = self._parse_css_length(val)
        
        # Check for align attribute (fallback if text-align not in style)
        if 'text_align' not in current_context and 'align' in element.attrs:
            current_context['text_align'] = element.attrs['align']

        # Handle hr elements first to avoid conflicts with heading detection
        if element.name == 'hr':
            self._add_horizontal_line()
        
        elif element.name.startswith('h') and len(element.name) == 2:
            # Ensure it's a valid heading (h1-h6) before converting to int
            try:
                level = int(element.name[1])
                if 1 <= level <= 6:
                    self._add_heading(element.text.strip(), level)
                # If level is not between 1-6, fall through to other conditions
            except ValueError:
                # Not a valid heading, fall through to other conditions
                pass
        
        elif element.name == 'p':
            if str(element).strip():
                # Check if this paragraph contains <br> tags
                has_br_tags = '<br>' in str(element) or '<br/>' in str(element)
                has_list_format = re.search(r'\(\d+\)\s+', str(element))
                
                if has_br_tags:
                    # Get the complete HTML string and split by <br> tags
                    html_content = str(element)
                    # Remove the outer <p> tags
                    html_content = re.sub(r'^<p[^>]*>|</p>$', '', html_content)
                    # Split by <br> tags
                    items_html = re.split(r'<br\s*/?>', html_content)
                    
                    for item_html in items_html:
                        if item_html.strip():
                            # Create a new paragraph for each line (hard line break)
                            p = self.doc.add_paragraph()
                            
                            # Set alignment if specified in context
                            if 'text_align' in current_context:
                                text_align = current_context['text_align'].lower()
                                if text_align == 'center':
                                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                elif text_align == 'right':
                                    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                                elif text_align == 'left':
                                    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                            
                            # Set text indent if specified in context
                            if 'text_indent' in current_context:
                                p.paragraph_format.first_line_indent = current_context['text_indent']
                            else:
                                # Default indentation for body text
                                p.paragraph_format.first_line_indent = Pt(21)
                            
                            # Set reduced space after for list items
                            if re.search(r'^\(\d+\)\s+', item_html):
                                p.paragraph_format.space_after = Pt(0)
                            
                            # Parse the item HTML and process inline elements
                            item_soup = BeautifulSoup(item_html, 'html.parser')
                            self._process_inline_elements(item_soup, p, current_context.get('font_name'), current_context.get('font_size'))
                else:
                    # Create a single paragraph
                    p = self.doc.add_paragraph()
                    
                    # Set alignment if specified in context
                    if 'text_align' in current_context:
                        text_align = current_context['text_align'].lower()
                        if text_align == 'center':
                            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        elif text_align == 'right':
                            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                        elif text_align == 'left':
                            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    
                    # Set text indent if specified in context
                    if 'text_indent' in current_context:
                        p.paragraph_format.first_line_indent = current_context['text_indent']
                    else:
                        # Default indentation for body text
                        p.paragraph_format.first_line_indent = Pt(21)
                    
                    # Check if it's a custom list item
                    if re.search(r'^\(\d+\)\s+', element.text.strip()):
                        p.paragraph_format.space_after = Pt(0)
                    
                    # Process all inline elements
                    self._process_inline_elements(element, p, current_context.get('font_name'), current_context.get('font_size'))
        

        
        elif element.name == 'table':
            self._add_table(element)
        
        elif element.name in ['ul', 'ol']:
            self._add_list(element, ordered=element.name == 'ol')
        
        elif element.name == 'pre':
            code = element.find('code')
            if code:
                p = self.doc.add_paragraph()
                run = p.add_run(code.get_text())
                
                # Try to apply template style for code blocks
                try:
                    p.style = "代码块"
                except Exception:
                    # Fallback to default style if template style not found and template not loaded
                    if not self.template_loaded:
                        self._apply_style(run, self.styles['code'])
        
        elif element.name == 'div':
            # Process all children elements recursively with updated context
            for child in element.children:
                self._process_element(child, current_context)
        
        else:
            # Process all children elements recursively
            for child in element.children:
                self._process_element(child, current_context)

    def _process_inline_elements(self, element, paragraph, font_name=None, font_size=None):
        """Process inline elements like bold, italic, etc."""
        # 逐节点处理元素内容，确保所有文本都被正确添加
        for child in element.contents:
            if isinstance(child, str):
                # 处理文本节点，清理多余换行符，保留原始空格（包括首行缩进的全角空格）
                text = child.replace('\n', '')
                if text.strip() or text.startswith('　'):
                    self._add_inline_content(paragraph, text, 'paragraph', False, False, font_name, font_size)
            elif child.name == 'br':
                # 处理换行符
                br_run = paragraph.add_run()
                br = OxmlElement('w:br')
                br_run._r.append(br)
            elif child.name == 'strong':
                # 处理加粗文本
                text = child.text.replace('\n', '').strip()
                if text:
                    self._add_inline_content(paragraph, text, 'strong', True, False, font_name, font_size)
            elif child.name == 'em':
                # 处理斜体文本
                text = child.text.replace('\n', '').strip()
                if text:
                    self._add_inline_content(paragraph, text, 'emphasis', False, True, font_name, font_size)
            elif child.name == 'code':
                # 处理代码
                text = child.text.replace('\n', '').strip()
                if text:
                    self._add_inline_content(paragraph, text, 'code', False, False, font_name, font_size)
            elif child.name == 'a':
                # 处理链接
                text = child.text.replace('\n', '').strip()
                if text:
                    self._add_inline_content(paragraph, text, 'paragraph', False, False, font_name, font_size)
                    # 链接样式可以在这里添加
            else:
                # 处理其他元素（递归）
                self._process_inline_elements(child, paragraph, font_name, font_size)

    def _add_table(self, table_element):
        """Add a table to the document with template styles."""
        rows = table_element.find_all('tr')
        if not rows:
            return
        
        num_cols = max(len(row.find_all(['td', 'th'])) for row in rows)
        table = self.doc.add_table(rows=0, cols=num_cols)
        table.style = 'Table Grid'
        table.allow_autofit = True
        
        # Process header
        header_cells = rows[0].find_all(['th', 'td'])
        header_row = table.add_row()
        for j, cell in enumerate(header_cells):
            table_cell = header_row.cells[j]
            paragraph = table_cell.paragraphs[0]
            run = paragraph.add_run(cell.get_text().strip())
            run.bold = True
            
            # Try to apply template style for table header
            try:
                paragraph.style = "表头"
            except Exception:
                # Fallback to default style if template style not found and template not loaded
                if not self.template_loaded:
                    self._apply_style(run, self.styles.get('table_header', self.styles['strong']))
                

        
        # Process rows
        for row in rows[1:]:
            cells = row.find_all(['td', 'th'])
            table_row = table.add_row()
            for j, cell in enumerate(cells):
                table_cell = table_row.cells[j]
                paragraph = table_cell.paragraphs[0]
                run = paragraph.add_run(cell.get_text().strip())
                
                # Try to apply template style for table content
                try:
                    paragraph.style = "表内"
                except Exception:
                    # Fallback to default style if template style not found and template not loaded
                    if not self.template_loaded:
                        self._apply_style(run, self.styles.get('table_cell', self.styles['paragraph']))
                        


    def _add_horizontal_line(self):
        """Add a page break."""
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()
        run.add_break(WD_BREAK.PAGE)

    def _add_list(self, list_element, ordered=False):
        """Add an ordered or unordered list as plain text to ensure stability."""
        # Find all direct 'li' children
        list_items = list_element.find_all('li', recursive=False)
        
        for index, item in enumerate(list_items):
            paragraph = self.doc.add_paragraph()
            level = len(item.find_parents(['ul', 'ol'])) - 1
            
            # Apply 'List Paragraph' style for font/spacing consistency
            try:
                paragraph.style = 'List Paragraph'
            except:
                paragraph.style = 'Normal'
            
            # Create the prefix (bullet or number)
            # We use plain text instead of Word auto-numbering to control start value and format
            prefix = ""
            if ordered:
                prefix = f"{index + 1}. "
            else:
                # Use standard bullet point
                prefix = "• "
            
            # Add prefix manually
            run_prefix = paragraph.add_run(prefix)
            # You might want to make the prefix font match the style?
            # self._apply_style(run_prefix, self.styles['paragraph'])
            
            # 检查列表项内部是否包含p标签
            p_tag = item.find('p')
            if p_tag:
                # 如果包含p标签，处理内容
                self._process_inline_elements(p_tag, paragraph)
            else:
                # 否则处理列表项本身
                self._process_inline_elements(item, paragraph)


def main():
    """
    Main function for command-line interface
    """
    import argparse
    
    parser = argparse.ArgumentParser(description='Markdown Converter - Convert Markdown to DOCX, HTML, and TXT formats')
    parser.add_argument('input_file', help='Path to the input Markdown file')
    parser.add_argument('format', choices=['docx', 'html', 'txt'], help='Output format')
    parser.add_argument('output_file', nargs='?', help='Path to the output file (optional)')
    
    args = parser.parse_args()
    
    try:
        if args.format == 'docx':
            result = convert_markdown_to_docx(args.input_file, args.output_file)
            print(f"✓ {result}")
        elif args.format == 'html':
            result = convert_markdown_to_html(args.input_file, args.output_file)
            print(f"✓ {result}")
        elif args.format == 'txt':
            result = convert_markdown_to_plain_text(args.input_file, args.output_file)
            print(f"✓ {result}")
        else:
            print(f"Error: Unsupported format '{args.format}'")
            return 1
        
        return 0
    except Exception as e:
        print(f"✗ Error: {e}")
        return 1


if __name__ == "__main__":
    main()
